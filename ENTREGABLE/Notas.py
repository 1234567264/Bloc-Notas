# -------- IMPORTS --------
import tkinter as tk
from tkinter import messagebox, ttk, filedialog
from datetime import datetime
from tkcalendar import Calendar
import sqlite3
import os
from docx import Document
from fpdf import FPDF

# -------- BASE DE DATOS --------
conn = sqlite3.connect('Mis_Tareas.db')
cursor = conn.cursor()
cursor.execute("""
    CREATE TABLE IF NOT EXISTS tareas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        titulo TEXT,
        descripcion TEXT,
        prioridad TEXT,
        fecha_limite TEXT,
        fecha_creacion TEXT,
        historial_actualizacion TEXT,
        carpeta_guardado TEXT
    )
""")
conn.commit()


# -------- CLASE PRINCIPAL --------
class Aplicacion:
    def __init__(self, root):
        self.bg_color = "#fef5c3"
        self.title_font = ("Georgia", 14, "bold", "underline")
        self.content_font = ("Calibri", 11)
        self.highlight = "#4CAF50"
        self.fg_color = "#333333"

        self.root = root
        self.root.title("🧠 Gestor de Tareas Empresarial")
        self.root.configure(bg=self.bg_color)
        self.root.geometry("950x550")
        self.root.resizable(True, True)

        style = ttk.Style()

        # 🟦 Frame superior para búsqueda + botones
        top_frame = tk.Frame(self.root, bg=self.bg_color)
        top_frame.pack(fill="x", padx=20, pady=(10, 5))

        # 🔎 Entrada de búsqueda tipo Google
        self.entrada_busqueda = tk.Entry(top_frame, font=("Segoe UI", 10), width=40)
        self.entrada_busqueda.pack(side="left", padx=(0, 5), ipady=3)
        self.entrada_busqueda.insert(0, "Buscar por título...")
        self.entrada_busqueda.bind("<FocusIn>", lambda e: self.entrada_busqueda.delete(0, tk.END))
        self.entrada_busqueda.bind("<KeyRelease>", lambda e: self.buscar_tareas())

        # 🔍 Botón de búsqueda
        tk.Button(top_frame, text="🔍", bg=self.highlight, fg="white", cursor="hand2",
                relief="flat", font=("Segoe UI", 10),
                command=self.buscar_tareas).pack(side="left", padx=(0, 20))

        # -------- LISTA DE TAREAS (Treeview) --------
        style.configure("Treeview", rowheight=30, font=self.content_font)

        self.lista_tareas = ttk.Treeview(self.root, columns=("ID", "Titulo", "Prioridad", "Fecha limite"), show='headings', height=10)

        for col in ("ID", "Titulo", "Prioridad", "Fecha limite"):
            self.lista_tareas.heading(col, text=col)

        # Tamaño y alineación de columnas
        self.lista_tareas.column("ID", width=50, anchor="center")
        self.lista_tareas.column("Titulo", width=350, anchor="w")
        self.lista_tareas.column("Prioridad", width=100, anchor="center")
        self.lista_tareas.column("Fecha limite", width=150, anchor="center")

        self.lista_tareas.pack(pady=20, padx=10, fill="x")

        # Scrollbar vertical
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=self.lista_tareas.yview)
        scrollbar.pack(side="right", fill="y")
        self.lista_tareas.configure(yscrollcommand=scrollbar.set)

        # -------- BOTONES --------
        btn_frame = tk.Frame(self.root, bg=self.bg_color)
        btn_frame.pack(pady=10)

        botones = [
            ("➕ Agregar", self.agregar_tarea),
            ("✏️ Actualizar", self.verificar_seleccion(self.actualizar_tarea)),
            ("🗑️ Eliminar", self.verificar_seleccion(self.eliminar_tarea)),
            ("📖 Ver Contenido", self.verificar_seleccion(self.ver_contenido)),
            ("🔎 Detalles", self.verificar_seleccion(self.ver_detalles))
        ]

        for texto, comando in botones:
            tk.Button(btn_frame, text=texto, width=15, font=("Segoe UI", 10, "bold"),
                    bg=self.highlight, fg='white', activebackground="#388e3c",
                    relief="flat", command=comando, cursor="hand2").pack(side=tk.LEFT, padx=6)

        self.actualizar_lista()

    
    def buscar_tareas(self):
        query = self.entrada_busqueda.get().strip().lower()
        for i in self.lista_tareas.get_children():
            self.lista_tareas.delete(i)

        if not query:
            cursor.execute("SELECT rowid, * FROM tareas")
            filas = cursor.fetchall()
        else:
            cursor.execute("""
                SELECT rowid, * FROM tareas 
                WHERE LOWER(titulo) LIKE ? OR LOWER(prioridad) LIKE ?
            """, ('%' + query + '%', '%' + query + '%'))
            filas = cursor.fetchall()

        for fila in filas:
            # fila[0] = rowid, fila[1] = título, fila[3] = prioridad, fila[4] = fecha límite
            self.lista_tareas.insert("", tk.END, values=(fila[0], fila[2], fila[4], fila[5]))
        
        print(fila)  # te muestra: (rowid, titulo, descripcion, prioridad, fecha_limite, fecha_creacion, historial, carpeta_guardado)




    def verificar_seleccion(self, funcion):
        def wrapper():
            seleccion = self.lista_tareas.selection()
            if not seleccion:
                messagebox.showwarning("Aviso", "Debes seleccionar una tarea para realizar esta acción.")
                return
            funcion()
        return wrapper

    def actualizar_lista(self):
        for i in self.lista_tareas.get_children():
            self.lista_tareas.delete(i)
        cursor.execute("SELECT rowid, titulo, prioridad, fecha_limite FROM tareas")
        for fila in cursor.fetchall():
            self.lista_tareas.insert('', 'end', values=(fila[0], fila[1], fila[2], fila[3]))
    def eliminar_tarea(self):
        seleccion = self.lista_tareas.selection()
        if not seleccion:
            messagebox.showwarning("Sin selección", "Debes seleccionar una tarea para eliminar.")
            return

        id_tarea = self.lista_tareas.item(seleccion[0])['values'][0]

        confirmacion = messagebox.askyesno("¿Eliminar tarea?", "¿Estás seguro de que deseas eliminar esta tarea?")
        if not confirmacion:
            return

        # Eliminar archivo .txt correspondiente si existe
        cursor.execute("SELECT carpeta_guardado, titulo FROM tareas WHERE rowid = ?", (id_tarea,))
        resultado = cursor.fetchone()
        if resultado:
            carpeta_guardado, titulo = resultado
            archivo_txt = os.path.join(carpeta_guardado, f"{titulo}.txt")
            if os.path.exists(archivo_txt):
                try:
                    os.remove(archivo_txt)
                except Exception as e:
                    messagebox.showerror("Error al eliminar archivo", f"No se pudo eliminar el archivo:\n{archivo_txt}\n\n{e}")

        # Eliminar de la base de datos
        cursor.execute("DELETE FROM tareas WHERE rowid = ?", (id_tarea,))
        conn.commit()

        self.actualizar_lista()
        messagebox.showinfo("Eliminado", "Tarea eliminada correctamente.")



    def agregar_tarea(self):
        ventana = tk.Toplevel(self.root)
        ventana.title("➕ Nueva Tarea")
        ventana.geometry("500x580")
        ventana.configure(bg=self.bg_color)
        ventana.attributes("-topmost", True)  # Siempre al frente

        tk.Label(ventana, text="Título:", font=self.title_font, bg=self.bg_color).pack(pady=(10, 0))
        titulo_entry = tk.Entry(ventana, font=self.content_font, width=40)
        titulo_entry.pack(pady=5)

        tk.Label(ventana, text="Descripción (opcional):", font=self.title_font, bg=self.bg_color).pack(pady=(10, 0))
        descripcion_text = tk.Text(ventana, height=5, width=50, font=self.content_font, undo=True)
        descripcion_text.pack(pady=5)
        descripcion_text.focus_set()

        tk.Label(ventana, text="Prioridad:", font=self.title_font, bg=self.bg_color).pack(pady=(10, 0))
        prioridad_combo = ttk.Combobox(ventana, values=["Alta", "Media", "Baja"], state="readonly")
        prioridad_combo.pack(pady=5)
        prioridad_combo.set("Media")

        tk.Label(ventana, text="Fecha límite:", font=self.title_font, bg=self.bg_color).pack(pady=(10, 0))
        fecha_limite = tk.StringVar()
        entry_fecha = tk.Entry(ventana, textvariable=fecha_limite, state="readonly", font=self.content_font)
        entry_fecha.pack(pady=5)

        def abrir_calendario():
            top = tk.Toplevel(ventana)
            top.title("Seleccionar Fecha")
            top.attributes("-topmost", True)

            # Posición al costado del botón
            x = ventana.winfo_rootx() + 350
            y = ventana.winfo_rooty() + 200
            top.geometry(f"+{x}+{y}")

            cal = Calendar(top, selectmode='day', date_pattern='yyyy-mm-dd')
            cal.pack(expand=True, fill="both", padx=10, pady=10)

            def seleccionar_fecha():
                fecha_limite.set(cal.get_date())
                top.destroy()

            tk.Button(top, text="✅ OK", bg="black", fg="white", font=("Segoe UI", 10),
                    command=seleccionar_fecha, cursor="hand2").pack(pady=(0, 10))

        tk.Button(ventana, text="📅", font=("Segoe UI", 11),
                command=abrir_calendario, bg="#fff", cursor="hand2").pack()

        def guardar():
            titulo = titulo_entry.get().strip().title()
            descripcion = descripcion_text.get("1.0", tk.END).strip()
            prioridad = prioridad_combo.get()
            fecha = fecha_limite.get()
            fecha_creacion = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            if not titulo or not prioridad or not fecha:
                messagebox.showwarning("Campos vacíos", "Debes completar: título, prioridad y fecha límite.")
                return

            # Formateo de descripción si existe
            if descripcion:
                lineas = descripcion.splitlines()
                descripcion_formateada = ""
                for linea in lineas:
                    linea = linea.strip()
                    if not linea:
                        #descripcion_formateada += "\n\n"
                        continue
                    if not linea.startswith("-"):
                        linea = "- " + linea
                    if not linea.endswith("."):
                        linea = linea.rstrip('.') + "."
                    linea = linea[0:2] + linea[2:].capitalize()
                    descripcion_formateada += linea + "\n\n"
                descripcion_formateada = descripcion_formateada.strip()
            else:
                descripcion_formateada = ""

            historial = f"[{fecha_creacion}] Creación de la tarea."

            # Selección de carpeta superpuesta
            self.root.attributes("-topmost", False)  # Desactiva topmost en root para evitar conflicto
            ventana.attributes("-topmost", False)
            carpeta = filedialog.askdirectory(title="Selecciona carpeta donde guardar")
            ventana.attributes("-topmost", True)
            self.root.attributes("-topmost", True)

            if not carpeta:
                return

            cursor.execute("""INSERT INTO tareas 
                (titulo, descripcion, prioridad, fecha_limite, fecha_creacion, historial_actualizacion, carpeta_guardado)
                VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (titulo, descripcion_formateada, prioridad, fecha, fecha_creacion, historial, carpeta)
            )
            conn.commit()

            self.actualizar_lista()
            ventana.destroy()
            messagebox.showinfo("Guardado", "Tarea guardada correctamente.")

        # Botón guardar
        tk.Button(ventana, text="💾 Guardar", font=("Segoe UI", 10, "bold"),
                bg=self.highlight, fg="white", command=guardar,
                relief="flat", cursor="hand2").pack(pady=15)

        # Atajos de teclado
        descripcion_text.bind("<Control-z>", lambda e: descripcion_text.edit_undo())
        descripcion_text.bind("<Control-y>", lambda e: descripcion_text.edit_redo())
        descripcion_text.bind("<Control-a>", lambda e: (descripcion_text.tag_add("sel", "1.0", "end"), "break"))
        descripcion_text.bind("<Control-c>", lambda e: descripcion_text.event_generate("<<Copy>>"))
        descripcion_text.bind("<Control-v>", lambda e: descripcion_text.event_generate("<<Paste>>"))
        descripcion_text.bind("<<Paste>>", lambda e: "break")  # Evita pegar doble

        ventana.bind("<Control-s>", lambda e: guardar())


    def ver_contenido(self):
        seleccion = self.lista_tareas.selection()
        if not seleccion:
            messagebox.showwarning("Sin selección", "Debes seleccionar una tarea para ver el contenido.")
            return

        id_tarea = self.lista_tareas.item(seleccion[0])['values'][0]
        cursor.execute("SELECT titulo, descripcion, prioridad, fecha_limite, fecha_creacion FROM tareas WHERE rowid = ?", (id_tarea,))
        titulo, descripcion, prioridad, fecha_limite, fecha_creacion = cursor.fetchone()

        ventana = tk.Toplevel(self.root)
        ventana.title("📝 Contenido de la Nota")
        ventana.geometry("720x560")
        ventana.configure(bg=self.bg_color)
        ventana.transient(self.root)
        ventana.grab_set()

        top_frame = tk.Frame(ventana, bg=self.bg_color)
        top_frame.pack(fill="x", padx=20, pady=(15, 5))

        titulo_label = tk.Label(top_frame, text=titulo, font=("Georgia", 14, "bold", "underline"),
                                bg=self.bg_color, fg="#000")
        titulo_label.pack(anchor="center")

        right_info = tk.Frame(top_frame, bg=self.bg_color)
        right_info.pack(anchor="ne", side="right")

        tk.Label(right_info, text=f"Creación: {fecha_creacion}", font=("Calibri", 9), bg=self.bg_color).pack(anchor="e")
        tk.Label(right_info, text=f"Límite: {fecha_limite}", font=("Calibri", 9), bg=self.bg_color).pack(anchor="e")

        color = "#e53935" if prioridad == "Alta" else "#fb8c00" if prioridad == "Media" else "#43a047"
        prioridad_box = tk.Label(right_info, text=f"Prioridad: {prioridad}", bg=color,
                                fg="white", font=("Calibri", 10, "bold"), padx=8, pady=2)
        prioridad_box.pack(anchor="e", pady=5)

        text_area = tk.Text(ventana, wrap="word", font=self.content_font, height=20,
                            bg="#fff", undo=True)
        text_area.insert(tk.END, descripcion)
        text_area.config(state="disabled")
        text_area.pack(padx=20, pady=(0, 10), expand=True, fill="both")

        boton_frame = tk.Frame(ventana, bg=self.bg_color)
        boton_frame.pack(pady=10)

        def desbloquear():
            text_area.config(state="normal")
            text_area.focus_set()

        def guardar_edicion():
            contenido = text_area.get("1.0", tk.END).strip()
            lineas = contenido.splitlines()
            nuevo_contenido = ""
            for linea in lineas:
                linea = linea.strip()
                if not linea:
                    continue  # Saltamos líneas vacías (no se duplican saltos)
                if not linea.startswith("-"):
                    linea = "- " + linea
                if not linea.endswith("."):
                    linea = linea.rstrip('.') + "."
                linea = linea[0:2] + linea[2:].capitalize()
                nuevo_contenido += linea + "\n\n"


            nuevo_contenido = nuevo_contenido.strip()

            fecha_edicion = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cursor.execute("SELECT historial_actualizacion FROM tareas WHERE rowid = ?", (id_tarea,))
            historial = cursor.fetchone()[0]
            historial += f"\n[{fecha_edicion}] Edición desde ver contenido."

            cursor.execute("UPDATE tareas SET descripcion = ?, historial_actualizacion = ? WHERE rowid = ?",
                        (nuevo_contenido, historial, id_tarea))
            conn.commit()

            text_area.delete("1.0", tk.END)
            text_area.insert(tk.END, nuevo_contenido)
            text_area.config(state="disabled")
            messagebox.showinfo("Guardado", "Cambios guardados correctamente.")
            ventana.destroy()
        def confirmar_cierre():
                if text_area.edit_modified():
                    respuesta = messagebox.askyesnocancel("Salir sin guardar", "¿Deseas guardar los cambios antes de salir?")
                    if respuesta is None:
                        return  # Cancelar salida
                    elif respuesta:
                        guardar_edicion()
                    else:
                        ventana.destroy()
                else:
                    ventana.destroy()

        ventana.protocol("WM_DELETE_WINDOW", confirmar_cierre)

        def exportar_pdf():
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            from reportlab.pdfbase import pdfmetrics

            pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))
            carpeta = filedialog.askdirectory(title="Selecciona carpeta para exportar")
            if not carpeta:
                return
            path = os.path.join(carpeta, f"{titulo}.pdf")
            c = canvas.Canvas(path, pagesize=letter)
            c.setFont("HeiseiMin-W3", 12)
            y = 750
            for linea in text_area.get("1.0", tk.END).splitlines():
                c.drawString(40, y, linea)
                y -= 20
                if y < 50:
                    c.showPage()
                    c.setFont("HeiseiMin-W3", 12)
                    y = 750
            c.save()
            messagebox.showinfo("Exportado", f"PDF guardado en:\n{path}")

        def exportar_word():
            from docx import Document
            carpeta = filedialog.askdirectory(title="Selecciona carpeta para exportar")
            if not carpeta:
                return
            path = os.path.join(carpeta, f"{titulo}.docx")
            contenido = text_area.get("1.0", tk.END).strip()

            try:
                doc = Document()
                doc.add_heading(titulo, level=1)
                for linea in contenido.split("\n"):
                    if linea.strip():
                        doc.add_paragraph(linea.strip())
                    else:
                        doc.add_paragraph()
                doc.save(path)
                messagebox.showinfo("Exportado", f"Word guardado en:\n{path}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo exportar a Word:\n{e}")

        def retroceder():
            try:
                text_area.edit_undo()
            except:
                pass

        def avanzar():
            try:
                text_area.edit_redo()
            except:
                pass

        tk.Button(boton_frame, text="🔓 Editar", bg=self.highlight, fg="white",
                font=("Segoe UI", 10), relief="flat", cursor="hand2",
                command=desbloquear).pack(side="left", padx=5)

        tk.Button(boton_frame, text="↩️ Retroceder", bg="#FFA726", fg="white",
                font=("Segoe UI", 10), relief="flat", cursor="hand2",
                command=retroceder).pack(side="left", padx=5)

        tk.Button(boton_frame, text="↪️ Avanzar", bg="#29B6F6", fg="white",
                font=("Segoe UI", 10), relief="flat", cursor="hand2",
                command=avanzar).pack(side="left", padx=5)

        tk.Button(boton_frame, text="💾 Guardar", bg=self.highlight, fg="white",
                font=("Segoe UI", 10), relief="flat", cursor="hand2",
                command=guardar_edicion).pack(side="left", padx=5)

        tk.Button(boton_frame, text="📄 Exportar PDF", bg="#2196F3", fg="white",
                font=("Segoe UI", 10), relief="flat", cursor="hand2",
                command=exportar_pdf).pack(side="left", padx=5)

        tk.Button(boton_frame, text="📝 Exportar Word", bg="#9C27B0", fg="white",
                font=("Segoe UI", 10), relief="flat", cursor="hand2",
                command=exportar_word).pack(side="left", padx=5)

        # Atajos de teclado corregidos
        def pegar_corregido(event=None):
            try:
                clipboard = ventana.clipboard_get()
                text_area.insert(tk.INSERT, clipboard)
            except:
                pass
            return "break"

        text_area.bind("<Control-z>", lambda e: retroceder())
        text_area.bind("<Control-y>", lambda e: avanzar())
        text_area.bind("<Control-a>", lambda e: (text_area.tag_add("sel", "1.0", "end"), "break"))
        text_area.bind("<Control-c>", lambda e: text_area.event_generate("<<Copy>>"))
        text_area.bind("<Control-v>", pegar_corregido)


    def ver_detalles(self):
        seleccion = self.lista_tareas.selection()
        if not seleccion:
            messagebox.showwarning("Sin selección", "Debes seleccionar una tarea para ver los detalles.")
            return

        id_tarea = self.lista_tareas.item(seleccion[0])['values'][0]
        cursor.execute("SELECT titulo, prioridad, fecha_limite, fecha_creacion, historial_actualizacion FROM tareas WHERE rowid = ?", (id_tarea,))
        titulo, prioridad, fecha_limite, fecha_creacion, historial = cursor.fetchone()

        ventana = tk.Toplevel(self.root)
        ventana.title("📋 Detalles de la Tarea")
        ventana.geometry("600x400")
        ventana.configure(bg=self.bg_color)
        ventana.transient(self.root)
        ventana.grab_set()

        # Frame superior
        top_frame = tk.Frame(ventana, bg=self.bg_color)
        top_frame.pack(fill="x", padx=20, pady=(15, 5))

        # Título
        titulo_label = tk.Label(top_frame, text=titulo, font=("Georgia", 14, "bold", "underline"),
                                bg=self.bg_color, fg="#000")
        titulo_label.pack(anchor="center")

        # Info lateral derecha
        right_info = tk.Frame(top_frame, bg=self.bg_color)
        right_info.pack(anchor="ne", side="right")

        tk.Label(right_info, text=f"Creación: {fecha_creacion}", font=("Calibri", 9), bg=self.bg_color).pack(anchor="e")
        tk.Label(right_info, text=f"Límite: {fecha_limite}", font=("Calibri", 9), bg=self.bg_color).pack(anchor="e")

        color = "#e53935" if prioridad == "Alta" else "#fb8c00" if prioridad == "Media" else "#43a047"
        prioridad_box = tk.Label(right_info, text=f"Prioridad: {prioridad}", bg=color,
                                fg="white", font=("Calibri", 10, "bold"), padx=8, pady=2)
        prioridad_box.pack(anchor="e", pady=5)

        # Historial
        historial_label = tk.Label(ventana, text="🕓 Historial de Actualizaciones", font=("Calibri", 11, "bold"),
                                bg=self.bg_color)
        historial_label.pack(pady=(10, 0))

        historial_text = tk.Text(ventana, wrap="word", font=("Consolas", 10), height=12,
                                bg="#fdfdfd", state="normal")
        historial_text.insert(tk.END, historial)
        historial_text.config(state="disabled")
        historial_text.pack(padx=20, pady=(5, 15), fill="both", expand=True)


    def actualizar_tarea(self):
        seleccion = self.lista_tareas.selection()
        if not seleccion:
            messagebox.showerror("Error", "Debes seleccionar una tarea para actualizar.")
            return

        id_tarea = self.lista_tareas.item(seleccion[0])['values'][0]
        cursor.execute("SELECT titulo, descripcion, prioridad, fecha_limite, fecha_creacion, carpeta_guardado FROM tareas WHERE rowid = ?", (id_tarea,))
        titulo, descripcion, prioridad, fecha_limite, fecha_creacion, carpeta_guardado = cursor.fetchone()

        ventana = tk.Toplevel(self.root)
        ventana.title("✏️ Actualizar Tarea")
        ventana.geometry("500x560")
        ventana.configure(bg=self.bg_color)
        ventana.attributes('-topmost', True)

        cambios_guardados = [False]

        def al_cerrar():
            if not cambios_guardados[0]:
                respuesta = messagebox.askyesno("¿Salir sin guardar?", "Tienes cambios sin guardar. ¿Deseas salir?", parent=ventana)
                if not respuesta:
                    return
            ventana.destroy()

        ventana.protocol("WM_DELETE_WINDOW", al_cerrar)

        tk.Label(ventana, text="Título:", bg=self.bg_color, font=("Segoe UI", 10, "bold")).pack(padx=20, anchor="w")
        entry_titulo = tk.Entry(ventana, font=("Segoe UI", 10))
        entry_titulo.insert(0, titulo)
        entry_titulo.pack(fill="x", padx=20, pady=(0, 10))

        tk.Label(ventana, text="Fecha límite:", bg=self.bg_color, font=("Segoe UI", 10, "bold")).pack(padx=20, anchor="w")
        frame_fecha = tk.Frame(ventana, bg=self.bg_color)
        frame_fecha.pack(fill="x", padx=20, pady=(0, 10))

        entry_fecha = tk.Entry(frame_fecha, font=("Segoe UI", 10))
        entry_fecha.insert(0, fecha_limite)
        entry_fecha.pack(side="left", fill="x", expand=True)

        def abrir_calendario():
            top = tk.Toplevel(ventana)
            top.title("Calendario")
            top.attributes("-topmost", True)
            top.configure(bg=self.bg_color)
            top.geometry("+%d+%d" % (ventana.winfo_rootx() + 250, ventana.winfo_rooty() + 180))
            cal = Calendar(top, selectmode='day', date_pattern='yyyy-mm-dd')
            cal.pack(padx=10, pady=10)

            def seleccionar_fecha():
                entry_fecha.delete(0, tk.END)
                entry_fecha.insert(0, cal.get_date())
                top.destroy()

            tk.Button(top, text="✅ OK", bg="black", fg="white", command=seleccionar_fecha).pack(pady=(0, 10))

        tk.Button(frame_fecha, text="📅", font=("Segoe UI", 10), bg="#f0f0f0", command=abrir_calendario, cursor="hand2").pack(side="left", padx=5)

        tk.Label(ventana, text="Prioridad:", bg=self.bg_color, font=("Segoe UI", 10, "bold")).pack(padx=20, anchor="w")
        combo_prioridad = ttk.Combobox(ventana, values=["Alta", "Media", "Baja"], state="readonly")
        combo_prioridad.set(prioridad)
        combo_prioridad.pack(fill="x", padx=20, pady=(0, 10))

        tk.Label(ventana, text="Descripción:", bg=self.bg_color, font=("Segoe UI", 10, "bold")).pack(padx=20, anchor="w")
        text_area = tk.Text(ventana, wrap="word", height=10, undo=True)
        text_area.insert("1.0", descripcion)
        text_area.pack(fill="both", expand=True, padx=20, pady=(0, 10))

        def guardar_cambios():
            nuevo_titulo = entry_titulo.get().strip().title()
            nueva_fecha = entry_fecha.get().strip()
            nueva_prioridad = combo_prioridad.get().strip()
            nuevo_contenido = text_area.get("1.0", tk.END).strip()

            if not nuevo_titulo or not nueva_fecha or not nueva_prioridad:
                messagebox.showerror("Error", "Por favor, completa todos los campos.")
                return

            lineas = nuevo_contenido.splitlines()
            contenido_formateado = "\n\n".join([
                linea if linea.startswith("- ") and linea.endswith(".")
                else f"- {linea.lstrip('- ').capitalize().rstrip('.') + '.'}" for linea in lineas if linea.strip()
            ])

            fecha_actual = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cursor.execute("SELECT historial_actualizacion FROM tareas WHERE rowid = ?", (id_tarea,))
            historial = cursor.fetchone()[0] or ""
            historial += f"\n[{fecha_actual}] Actualización de tarea."

            cursor.execute("""
                UPDATE tareas
                SET titulo = ?, descripcion = ?, prioridad = ?, fecha_limite = ?, historial_actualizacion = ?
                WHERE rowid = ?
            """, (nuevo_titulo, contenido_formateado, nueva_prioridad, nueva_fecha, historial, id_tarea))
            conn.commit()

            # Actualizar archivo TXT si existe
            if carpeta_guardado:
                try:
                    ruta_txt = os.path.join(carpeta_guardado, f"{nuevo_titulo}.txt")
                    with open(ruta_txt, "w", encoding="utf-8") as archivo:
                        archivo.write(f"{nuevo_titulo}\n\n{contenido_formateado}")
                except Exception as e:
                    messagebox.showwarning("Error", f"No se pudo actualizar el archivo:\n{e}")

            cambios_guardados[0] = True
            ventana.destroy()
            self.actualizar_lista()
            messagebox.showinfo("Éxito", "Tarea actualizada correctamente.")

        tk.Button(ventana, text="💾 Guardar Cambios", bg="#4caf50", fg="white", font=("Segoe UI", 10),
                relief="flat", command=guardar_cambios, cursor="hand2").pack(pady=10)

        # Atajos
        text_area.bind("<Control-z>", lambda e: text_area.edit_undo())
        text_area.bind("<Control-y>", lambda e: text_area.edit_redo())
        text_area.bind("<Control-a>", lambda e: (text_area.tag_add("sel", "1.0", "end"), "break"))
        text_area.bind("<Control-c>", lambda e: text_area.event_generate("<<Copy>>"))
        text_area.bind("<Control-v>", lambda e: text_area.event_generate("<<Paste>>"))
        text_area.bind("<Control-s>", lambda e: guardar_cambios())


    def formatear_contenido(self, texto):
        lineas = texto.splitlines()
        return "\n".join([
            linea if linea.startswith("- ") and linea.endswith(".")
            else f"- {linea.capitalize().rstrip('.') + '.'}" for linea in lineas if linea.strip()
        ])

# ---------- EJECUCIÓN ----------
if __name__ == "__main__":
    root = tk.Tk()
    app = Aplicacion(root)
    root.mainloop()
